from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture(
    'me.jpeg',
    width=Inches(2.0)
    )


name = input('what is your name?')
phone_number =input('what is your phone number?')
email = input('what is your emailbh?')

document.add_paragraph(
    name + ' ' + phone_number + ' ' + email
)


document.add_heading('About me')
document.add_paragraph( input('Tell me about your self?'))


# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter Company')
from_date = input("From Date")
to_date = input("To Date")

p.add_run(company + ' ').bold=True
p.add_run(from_date + '-'+ to_date + '\n').italic=True
experience_details = input(
    'Desribe your experience at ' + company)
p.add_run(experience_details)


# more experiences

while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ')
        from_date = input("From Date ")
        to_date = input("To Date ")

        p.add_run(company + ' ').bold=True
        p.add_run(from_date + '-'+ to_date + '\n').italic=True
        experience_details = input(
        'Desribe your experience at ' + company + ' ' )
        p.add_run(experience_details)
    else:
        break




document.save('cv.docx')